from docx import Document
from docx.shared import Pt
from pathlib import Path
import json


from pathlib import Path
from jinja2 import Environment, FileSystemLoader

BASE_DIR = Path(__file__).resolve().parent.parent   # 27-Delta/ 폴더
TEMPLATE_DIR = BASE_DIR / "templates"

env = Environment(loader=FileSystemLoader(str(TEMPLATE_DIR)))

#==========
# 공통 유틸
#==========

# def add_heading(doc, text, level=1):
#     h = doc.add_heading(text, level=level)
#     h.style.font.name = "Arial"
#     h.style.font.size = Pt(14 if level == 1 else 12)
#     return h

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    run = h.runs[0]
    run.font.name = "Arial"
    run.font.size = Pt(14 if level == 1 else 12)
    run.bold = True
    h.paragraph_format.space_before = Pt(50)
    h.paragraph_format.space_after = Pt(20)
    return h


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.style.font.name = "Arial"
    p.style.font.size = Pt(11)
    return p


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.name = "Arial"
    p.style.font.size = Pt(11)
    return p


#==========
# A
#==========

def export_A_to_docx(case_id: str):
    path = f"cache/{case_id}/export_A_full.json"
    data = json.load(open(path, encoding="utf-8"))

    exec_summary = data["executive_summary"]["executive_summary"]
    body = data["body"]

    doc = Document()

    add_heading(doc, f"Case Report – {case_id}", level=1)
    add_heading(doc, "Executive Summary", level=2)

    add_heading(doc, "What this case is about", level=3)
    add_para(doc, exec_summary["what_this_case_is_about"])

    add_heading(doc, "Key Points", level=3)
    for item in exec_summary["key_points_in_20_words"]:
        add_para(doc, f"- {item}")

    add_heading(doc, "Micro Takeaway", level=3)
    add_para(doc, exec_summary["micro_takeaway"])

    # -------- Metadata --------
    meta = body["metadata"]
    add_heading(doc, "Metadata", level=2)
    for k, v in meta.items():
        add_para(doc, f"{k}: {v}")

    # -------- Narrative --------
    nar = body["narrative"]

    add_heading(doc, "Narrative Summary", level=2)
    add_para(doc, "Fact Summary:\n" + nar["fact_summary"])

    add_para(doc, "\nPlaintiff Arguments:")
    for x in nar["plaintiff_arguments"]:
        add_para(doc, "- " + x)

    add_para(doc, "\nDefendant Arguments:")
    for x in nar["defendant_arguments"]:
        add_para(doc, "- " + x)

    add_para(doc, "\nLegal Context:")
    for x in nar["legal_context"]:
        add_para(doc, "- " + x)

    add_para(doc, "\nCourt Reasoning:")
    for x in nar["court_reasoning"]:
        add_para(doc, "- " + x)

    out_path = Path(f"cache/{case_id}/report_A_full.docx")
    doc.save(out_path)

    return str(out_path)


#==========
# B
#==========

def export_B_to_docx(case_id: str):
    path = f"cache/{case_id}/export_B_full.json"
    data = json.load(open(path, encoding="utf-8"))

    exec_summary = data["executive_summary"]["executive_summary"]
    body = data["body"]

    doc = Document()
    add_heading(doc, f"Case Report – {case_id}", level=1)

    # ---------------- Executive Summary ----------------
    add_heading(doc, "Executive Summary", level=2)

    add_heading(doc, "One Liner", level=3)
    add_para(doc, exec_summary["one_liner"])

    add_heading(doc, "Main Conflicts", level=3)
    for item in exec_summary["main_conflicts"]:
        add_para(doc, "- " + item)

    add_heading(doc, "Legal Direction", level=3)
    add_para(doc, exec_summary["legal_direction"])

    add_heading(doc, "Practical Implication", level=3)
    add_para(doc, exec_summary["practical_implication"])

    # ---------------- Metadata ----------------
    meta = body["metadata"]
    add_heading(doc, "Metadata", level=2)
    for k, v in meta.items():
        add_para(doc, f"{k}: {v}")

    # ---------------- Narrative ----------------
    nar = body["narrative"]
    add_heading(doc, "Narrative Summary", level=2)

    add_para(doc, "Fact Summary:\n" + nar["fact_summary"])
    add_para(doc, "\nPlaintiff Arguments:")
    for x in nar["plaintiff_arguments"]:
        add_para(doc, "- " + x)

    add_para(doc, "\nDefendant Arguments:")
    for x in nar["defendant_arguments"]:
        add_para(doc, "- " + x)

    add_para(doc, "\nLegal Context:")
    for x in nar["legal_context"]:
        add_para(doc, "- " + x)

    add_para(doc, "\nCourt Reasoning:")
    for x in nar["court_reasoning"]:
        add_para(doc, "- " + x)

    # ---------------- Issue Frame ----------------
    issue_groups = body["issue_frame"]["issue_groups"]

    add_heading(doc, "Issue Frame (Structured Reasoning)", level=2)

    for issue_name, content in issue_groups.items():
        add_heading(doc, f"Issue: {issue_name}", level=3)

        # 안전하게 get 처리
        p_args = content.get("plaintiff_arguments", [])
        d_args = content.get("defendant_arguments", [])
        laws = content.get("legal_context", [])
        reasoning = content.get("court_reasoning", [])

        add_para(doc, "\nPlaintiff Arguments:")
        for x in p_args:
            add_para(doc, "- " + x)

        add_para(doc, "\nDefendant Arguments:")
        for x in d_args:
            add_para(doc, "- " + x)

        add_para(doc, "\nLegal Context:")
        for x in laws:
            add_para(doc, "- " + x)

        add_para(doc, "\nCourt Reasoning:")
        for x in reasoning:
            add_para(doc, "- " + x)

        add_para(doc, "\n---\n")

    # ---------------- Statutes ----------------
    statutes = body["statutes"]
    add_heading(doc, "Referenced Statutes", level=2)

    for s in statutes:
        add_para(doc, f"- {s['title']} ({s['citation']})")

    out_path = Path(f"cache/{case_id}/report_B_full.docx")
    doc.save(out_path)
    return str(out_path)



#==========
# C
#==========

def export_C_to_docx(case_id: str):
    path = f"cache/{case_id}/export_C_full.json"
    data = json.load(open(path, encoding="utf-8"))

    exec_summary = data["executive_summary"]["executive_summary"]
    body = data["body"]

    doc = Document()

    add_heading(doc, f"Case Report – {case_id}", level=1)

    # --- Executive Summary ---
    add_heading(doc, "Executive Summary", level=2)
    add_para(doc, exec_summary["one_liner"])

    add_heading(doc, "Core Issues", level=3)
    for item in exec_summary["core_issues"]:
        add_para(doc, "- " + item)

    add_heading(doc, "Judicial Logic", level=3)
    add_para(doc, exec_summary["judicial_logic"]["how_the_court_thought"])

    add_heading(doc, "Legal Signals", level=3)
    for x in exec_summary["judicial_logic"].get("legal_context", []):
        add_para(doc, "- " + x)

    add_heading(doc, "Risk View – Taxpayer", level=3)
    add_para(doc, exec_summary["risk_view"]["taxpayer_risk"])

    add_heading(doc, "Risk View – Tax Authority", level=3)
    add_para(doc, exec_summary["risk_view"]["tax_authority_risk"])

    add_heading(doc, "Precedent Signal", level=3)
    add_para(doc, exec_summary["risk_view"]["precedent_signal"])



    # --- Appendix ---
    add_heading(doc, "Appendix", level=2)
    # --- Metadata / Narrative / Issue Logic / Statutes ---

    # ---------------- Metadata ----------------
    meta = body["metadata"]
    add_heading(doc, "Appendix A – Metadata", level=3)
    for k, v in meta.items():
        add_para(doc, f"{k}: {v}")


    # ---------------- Narrative ----------------
    nar = body["narrative"]
    add_heading(doc, "Appendix B – Narrative Summary", level=3)

    add_para(doc, "Fact Summary:\n" + nar["fact_summary"])
    add_para(doc, "\nPlaintiff Arguments:")
    for x in nar["plaintiff_arguments"]:
        add_para(doc, "- " + x)

    add_para(doc, "\nDefendant Arguments:")
    for x in nar["defendant_arguments"]:
        add_para(doc, "- " + x)

    add_para(doc, "\nLegal Context:")
    for x in nar["legal_context"]:
        add_para(doc, "- " + x)

    add_para(doc, "\nCourt Reasoning:")
    for x in nar["court_reasoning"]:
        add_para(doc, "- " + x)

    # ---------------- Issue Logic ----------------
    issue_logic = body["issue_logic"]

    add_heading(doc, "Appendix C – Issue Logic (Structured Reasoning)", level=3)

    # Global Outline
    if issue_logic.get("global_outline"):
        add_heading(doc, "Global Reasoning Outline", level=3)
        add_para(doc, issue_logic["global_outline"])

    # Main Issues
    if issue_logic.get("main_issues"):
        add_heading(doc, "Main Issues", level=3)
        for item in issue_logic["main_issues"]:
            add_para(doc, f"- {item}")

    # Detailed Chains
    chains = issue_logic.get("issue_logic_chains", [])
    for chain in chains:
        add_heading(doc, f"Issue: {chain.get('issue', 'Untitled Issue')}", level=3)

        # Premise
        if chain.get("premise"):
            add_para(doc, "Premise:")
            add_para(doc, chain["premise"])

        # Evidence
        if chain.get("evidence"):
            add_para(doc, "Evidence:")
            add_para(doc, chain["evidence"])

        # Rule
        if chain.get("rule"):
            add_para(doc, "Rule:")
            add_para(doc, chain["rule"])

        # Application
        if chain.get("application"):
            add_para(doc, "Application:")
            add_para(doc, chain["application"])

        # Inference
        if chain.get("inference"):
            add_para(doc, "Inference:")
            add_para(doc, chain["inference"])

        # Mini conclusion
        if chain.get("mini_conclusion"):
            add_para(doc, "Conclusion:")
            add_para(doc, chain["mini_conclusion"])

        add_para(doc, "\n---\n")

    # ---------------- Statutes ----------------
    statutes = body["statutes"]
    add_heading(doc, "Appendix D – Referenced Statutes", level=3)
    
    for s in statutes:
        add_para(doc, f"- {s['title']} ({s['citation']})")

    out_path = Path(f"cache/{case_id}/report_C_full.docx")
    doc.save(out_path)
    return str(out_path)
